import json
from docx import Document
from docxtpl import DocxTemplate


def read_docx(docx_path, items_per_subarray=142):
    doc = Document(docx_path)
    all_text_list = []

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    all_text_list.append(text)

    divided_data = [all_text_list[i:i + items_per_subarray] for i in range(0, len(all_text_list), items_per_subarray)]

    return divided_data

def create_json(data, json_filename):
    json_data = []
    for item in data:
        json_item = {
            "napravlen": item[2],
            "studentA": item[26],
            "title": item[29],
            "nauchruk": item[37],
            "role": item[41],
            "student": item[58],
            "studentU": item[115],
            "kval": item[121],
            "spec": item[126]
        }
        json_data.append(json_item)

    with open(json_filename, 'w', encoding='utf-8') as json_file:
        json.dump(json_data, json_file, ensure_ascii=False, indent=4)

docx_path = 'exprot2.docx'
all_text_data = read_docx(docx_path)

json_filename = 'output.json'
create_json(all_text_data, json_filename)

# Вставка данных в шаблон
def create_draft(data, output_docx_path):
    doc = DocxTemplate('shablon.docx')
    
    combined_doc = Document()
    
    for item in data:
        context = {
            'napravlenie': item.get('napravlen', ''),
            'studentA': item.get('studentA', ''),
            'title': item.get('title', ''),
            'nauchruc': item.get('nauchruk', ''),
            'rang': item.get('role', ''),
            'student': item.get('student', ''),
            'studentU': item.get('studentU', ''),
            'kval': item.get('kval', ''),
            'spec': item.get('spec', ''),
            'predgos': 'Баранчиков Алексей Иванович',
            'score': 'отлично',
        }
        
        if context['score'] == 'удовлетворительно':
            context['haracteransver1'] = 'Студент показал достаточный уровень подготовки. Недостаточно глубоко изучил и' 
            context['haracteransver2'] = 'проанализировал предметную область. При защите ВКР студент проявил неуверенность,'
            context['haracteransver3'] = 'показал слабое знание вопросов темы, не дал полного аргументированного ответа на'
            context['haracteransver4'] = 'заданные вопросы.'
        elif context['score'] == 'отлично':
            context['haracteransver1'] = 'Студент показал высокий уровень подготовки и глубокие системные знания,' 
            context['haracteransver2'] = 'свободно оперирует данными исследования, дал развернутые и полные ответы на'
            context['haracteransver3'] = 'поставленные вопросы'
        elif context['score'] == 'хорошо':
            context['haracteransver1'] = 'Студент показал высокий уровень подготовки и глубокие системные знания,' 
            context['haracteransver2'] = 'но на дополнительные вопросы комиссии были даны неполные ответы.'
        
        context['ekzscore'] = 'не предусмотрен учебным планом'
        context['nodata'] = '–'
        context['data'] = '14.06.2023'
        context['scorediplom'] = 'без отличия'
        context['predgossokr'] = 'Баранчиков А.И.'
        context['sekretgossokr'] = 'Трохаченкова Н.Н.'

        doc.render(context)

        for element in doc.element.body:
            combined_doc.element.body.append(element)

    combined_doc.save(output_docx_path)

json_file_path = 'output.json'
with open(json_file_path, 'r', encoding='utf-8') as json_file:
    all_data = json.load(json_file)
output_path = 'combined_generated_docx.docx'
create_draft(all_data, output_path)